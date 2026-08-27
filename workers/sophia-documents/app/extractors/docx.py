from ..schemas import ExtractedBlock, ExtractionResult


def extract_docx(buffer: bytes) -> ExtractionResult:
    try:
        from docx import Document
        from io import BytesIO
        document = Document(BytesIO(buffer))
    except Exception as exc:
        return ExtractionResult(status="erro", warnings=[f"DOCX nao pode ser lido: {exc}"])

    blocks: list[ExtractedBlock] = []
    index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            kind = "heading" if paragraph.style and paragraph.style.name.lower().startswith("heading") else "paragraph"
            blocks.append(ExtractedBlock(text=text, kind=kind, order_index=index, heading=text if kind == "heading" else None))
            index += 1
    for table in document.tables:
        rows = [[cell.text.strip().replace("|", "\\|") for cell in row.cells] for row in table.rows]
        if rows:
            header = "| " + " | ".join(rows[0]) + " |"
            separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
            body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
            blocks.append(ExtractedBlock(text="\n".join([header, separator, *body]), kind="table", order_index=index))
            index += 1
    return ExtractionResult(blocks=blocks, method="native")
